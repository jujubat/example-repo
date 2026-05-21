from flask import Blueprint, jsonify, request
from flask_jwt_extended import jwt_required, get_jwt_identity
from ..models import SOPDocument, db
from datetime import datetime
import os
import tempfile
from werkzeug.utils import secure_filename

admin_bp = Blueprint('admin', __name__)


@admin_bp.route('/')
def index():
    return jsonify({'msg': 'Admin console placeholder'})


@admin_bp.route('/sop', methods=['POST'])
@jwt_required()
def create_sop():
    """Create or update an SOP document with file upload support"""
    try:
        # Check if it's a file upload or manual entry
        if 'file' in request.files:
            # Handle file upload
            file = request.files['file']
            if file.filename == '':
                return jsonify({'error': 'No file selected'}), 400

            # Validate file type
            allowed_extensions = {'pdf', 'xlsx', 'xls', 'jpg', 'jpeg', 'png', 'txt', 'docx'}
            if not _allowed_file(file.filename, allowed_extensions):
                return jsonify({'error': 'Unsupported file type'}), 400

            # Save file temporarily and process it
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                file.save(temp_file.name)
                content = process_sop_file(temp_file.name, file.filename)
                os.unlink(temp_file.name)  # Clean up temp file

            if content is None:
                return jsonify({'error': 'Failed to process file'}), 400

        elif 'gsheet_url' in request.form:
            # Handle Google Sheets URL
            gsheet_url = request.form.get('gsheet_url')
            if not gsheet_url:
                return jsonify({'error': 'Google Sheets URL is required'}), 400

            content = _process_google_sheet(gsheet_url)

        else:
            # Handle manual entry
            content = request.form.get('content')
            if not content:
                return jsonify({'error': 'Content is required'}), 400

        # Get other form data
        category = request.form.get('category')
        title = request.form.get('title')
        summary = request.form.get('summary', '')

        if not category or not title:
            return jsonify({'error': 'Category and title are required'}), 400

        # Check if SOP exists for this category
        existing_sop = SOPDocument.query.filter_by(
            category=category,
            is_active=True
        ).first()

        if existing_sop:
            # Increment version and deactivate old one
            existing_sop.is_active = False
            new_version = existing_sop.version + 1
        else:
            new_version = 1

        # Create new SOP
        new_sop = SOPDocument(
            category=category,
            title=title,
            content=content,
            summary=summary,
            version=new_version,
            is_active=True,
            created_by=get_jwt_identity(),
            created_at=datetime.utcnow()
        )

        db.session.add(new_sop)
        db.session.commit()

        return jsonify({
            'message': 'SOP created successfully',
            'sop_id': new_sop.id,
            'version': new_version
        }), 201

    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to create SOP: {str(e)}'}), 500


def _allowed_file(filename, allowed_extensions):
    """Check if file has allowed extension"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in allowed_extensions


def process_sop_file(file_path, filename):
    """Process different file types and extract text content"""
    file_extension = filename.rsplit('.', 1)[1].lower()

    try:
        if file_extension == 'txt':
            return _process_txt_file(file_path)
        elif file_extension == 'pdf':
            return _process_pdf_file(file_path)
        elif file_extension in ['xlsx', 'xls']:
            return _process_excel_file(file_path)
        elif file_extension in ['jpg', 'jpeg', 'png']:
            return _process_image_file(file_path)
        elif file_extension == 'docx':
            return _process_docx_file(file_path)
        elif file_extension == 'gsheet':
            return _process_google_sheet(file_path)
        else:
            return None
    except Exception as e:
        print(f"Error processing file {filename}: {str(e)}")
        return None


def _process_txt_file(file_path):
    """Process text file"""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def _process_pdf_file(file_path):
    """Process PDF file"""
    try:
        import PyPDF2
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
    except ImportError:
        return "PDF processing not available. Please install PyPDF2."
    except Exception as e:
        return f"Error processing PDF: {str(e)}"


def _process_excel_file(file_path):
    """Process Excel file"""
    try:
        import pandas as pd
        df = pd.read_excel(file_path)
        # Convert all data to string and join
        text_content = []
        for col in df.columns:
            text_content.append(f"{col}:")
            for value in df[col].dropna():
                text_content.append(f"  {str(value)}")
            text_content.append("")
        return "\n".join(text_content)
    except ImportError:
        return "Excel processing not available. Please install pandas and openpyxl."
    except Exception as e:
        return f"Error processing Excel file: {str(e)}"


def _process_image_file(file_path):
    """Process image file using OCR"""
    try:
        import pytesseract
        from PIL import Image
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return text.strip()
    except ImportError:
        return "OCR processing not available. Please install pytesseract and pillow."
    except Exception as e:
        return f"Error processing image: {str(e)}"


def _process_docx_file(file_path):
    """Process Word document"""
    try:
        from docx import Document
        doc = Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        return "\n".join(text)
    except ImportError:
        return "Word document processing not available. Please install python-docx."
    except Exception as e:
        return f"Error processing Word document: {str(e)}"


def _process_google_sheet(sheet_url_or_id):
    """Process Google Sheet - basic implementation"""
    try:
        # Extract sheet ID from URL if it's a full URL
        if 'docs.google.com/spreadsheets' in sheet_url_or_id:
            # Extract ID from URL like: https://docs.google.com/spreadsheets/d/SPREADSHEET_ID/edit
            import re
            match = re.search(r'/spreadsheets/d/([a-zA-Z0-9-_]+)', sheet_url_or_id)
            if match:
                sheet_id = match.group(1)
            else:
                return "Invalid Google Sheets URL format"
        else:
            # Assume it's already a sheet ID
            sheet_id = sheet_url_or_id

        # For now, return a placeholder message
        # Full implementation would require Google Sheets API setup
        return f"Google Sheet detected (ID: {sheet_id}). To enable automatic processing, please configure Google Sheets API with proper authentication. For now, please copy the sheet content manually or contact administrator to set up the integration."

    except Exception as e:
        return f"Error processing Google Sheets URL: {str(e)}"


@admin_bp.route('/sop', methods=['GET'])
@jwt_required()
def get_sops():
    """Get all SOP documents"""
    sops = SOPDocument.query.order_by(
        SOPDocument.category,
        SOPDocument.version.desc()
    ).all()

    result = []
    for sop in sops:
        result.append({
            'id': sop.id,
            'category': sop.category,
            'title': sop.title,
            'content': sop.content,
            'summary': sop.summary,
            'version': sop.version,
            'is_active': sop.is_active,
            'created_by': sop.created_by,
            'created_at': sop.created_at.isoformat()
        })

    return jsonify(result)


@admin_bp.route('/sop/<int:sop_id>', methods=['PUT'])
@jwt_required()
def update_sop(sop_id):
    """Update SOP status (activate/deactivate)"""
    sop = SOPDocument.query.get_or_404(sop_id)
    data = request.get_json()

    if 'is_active' in data:
        sop.is_active = data['is_active']

    db.session.commit()

    return jsonify({'message': 'SOP updated successfully'})
